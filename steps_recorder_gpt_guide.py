import os
import re
import zipfile
from io import BytesIO

import openai
from bs4 import BeautifulSoup
from docx import Document

def process_steps_recorder_file(file_path):
    with zipfile.ZipFile(file_path, 'r') as z:
        mht_content = z.read('Steps/Default.mht').decode('utf-8')
        img_data = {os.path.basename(img_path): BytesIO(z.read(img_path)) for img_path in z.namelist() if img_path.startswith('Steps/')}
    
    soup = BeautifulSoup(mht_content, 'html.parser')
    step_elements = soup.select('.StepOuterDiv')
    
    steps = []
    for step_element in step_elements:
        step_text = step_element.select_one('.StepBullet').text.strip()
        step_img = step_element.select_one('.StepImg')['src'][6:]
        steps.append((step_text, step_img))
    
    return steps, img_data

def filter_steps(steps):
    filtered_steps = []
    for step in steps:
        filtered_step = re.sub(r'(Program:.*|UI Elements:.*)', '', step)
        filtered_step = re.sub(r'\s+', ' ', filtered_step).strip()
        filtered_steps.append(filtered_step)
    return filtered_steps

def get_gpt_guide(steps):
    prompt = "Create a guide based on the steps, but only include the critical steps because the guide will be read by technical users. Steps:\n"
    for i, step in enumerate(steps):
        prompt += f"{i + 1}. {step}\n"
    
    response = openai.Completion.create(
        engine="text-davinci-002",
        prompt=prompt,
        max_tokens=150,
        n=1,
        stop=None,
        temperature=0.5,
    )

    result_text = response.choices[0].text.strip()
    return result_text.split("\n")

def create_word_doc(guide, steps, img_data):
    doc = Document()

    for step in guide:
        step_number = int(re.search(r'\d+', step).group()) - 1
        step_text, step_img = steps[step_number]

        doc.add_paragraph().add_run(step_text).bold = True
        doc.add_picture(img_data[step_img], width=doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin)

    doc.save('guide.docx')

def main():
    # Configure the path to the Steps Recorder output file
    steps_recorder_file = r'E:\Python Stuff\Steps Recorder\steps_recorder_file.zip'

    steps, img_data = process_steps_recorder_file(steps_recorder_file)
    filtered_steps = filter_steps([step for step, _ in steps])
    guide = get_gpt_guide(filtered_steps)
    create_word_doc(guide, steps, img_data)

if __name__ == "__main__":
    main()
